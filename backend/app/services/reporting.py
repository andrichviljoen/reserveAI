from __future__ import annotations

import base64
import io
import json
from typing import Any

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches
from openai import OpenAI


def generate_ai_text(prompt: str, api_key: str | None) -> str:
    if not api_key:
        return "AI response unavailable because no OpenAI API key was provided."
    client = OpenAI(api_key=api_key)
    resp = client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": prompt}])
    return resp.choices[0].message.content or ""


def chat_completion(context_data: dict[str, Any], notes: str, messages: list[dict[str, str]], api_key: str | None) -> str:
    if not api_key:
        return "Please provide an OpenAI API key to use the A(I)ctuary chat."
    client = OpenAI(api_key=api_key)
    sys_msg = f"Senior Actuary Consultant. Context: {json.dumps(context_data)}. Notes: {notes}"
    resp = client.chat.completions.create(model="gpt-4o", messages=[{"role": "system", "content": sys_msg}] + messages)
    return resp.choices[0].message.content or ""


def _plot_to_buf(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_word(context_data: dict[str, Any], notes: str, api_key: str | None):
    prompt = (
        "Write an exhaustive actuarial reserve report. Include executive summary, outlier analysis, "
        "exclusion suggestions, AvE narrative, and bootstrap uncertainty discussion. "
        f"Context: {json.dumps(context_data)}. Notes: {notes}"
    )
    narrative = generate_ai_text(prompt, api_key)

    doc = Document()
    doc.add_heading("Exhaustive Actuarial Reserve Report", 0)
    doc.add_paragraph(narrative)

    tri_df = pd.DataFrame(context_data.get("tables", {}).get("triangle", []))
    if not tri_df.empty:
        doc.add_heading("1. Data & Development Patterns", level=1)
        fig, ax = plt.subplots(figsize=(10, 4))
        for col in [c for c in tri_df.columns if c not in ["origin", "development"]][:5]:
            ax.plot(tri_df.index, tri_df[col], alpha=0.8)
        ax.set_title("Cumulative Development")
        doc.add_picture(_plot_to_buf(fig), width=Inches(5.8))

    ave_df = pd.DataFrame(context_data.get("tables", {}).get("ave", []))
    if not ave_df.empty and all(c in ave_df.columns for c in ["Actual", "Expected"]):
        doc.add_heading("2. Actual vs Expected (AvE)", level=1)
        fig, ax = plt.subplots(figsize=(10, 4))
        plot_df = ave_df.iloc[::-1]
        y = range(len(plot_df))
        ax.hlines(y=y, xmin=plot_df["Actual"], xmax=plot_df["Expected"], color="grey", alpha=0.4)
        ax.scatter(plot_df["Actual"], y, label="Actual", color="#1f77b4")
        ax.scatter(plot_df["Expected"], y, label="Expected", color="#ff7f0e")
        ax.legend()
        doc.add_picture(_plot_to_buf(fig), width=Inches(5.8))

    buf = io.BytesIO()
    doc.save(buf)
    return "Exhaustive_Actuarial_Report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", base64.b64encode(buf.getvalue()).decode()


def generate_excel(context_data: dict[str, Any], notes: str, api_key: str | None):
    comment_prompt = (
        "Provide professional actuarial workbook commentary in four paragraphs: "
        "Data, Model, AvE, Risk. "
        f"Context: {json.dumps(context_data)}. Notes: {notes}"
    )
    commentary = generate_ai_text(comment_prompt, api_key).split("\n\n")
    output = io.BytesIO()

    tri_df = pd.DataFrame(context_data.get("tables", {}).get("triangle", []))
    ult_df = pd.DataFrame(context_data.get("tables", {}).get("ultimate", []))
    ibnr_df = pd.DataFrame(context_data.get("tables", {}).get("ibnr", []))
    ave_df = pd.DataFrame(context_data.get("tables", {}).get("ave", []))

    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        wb = writer.book
        h_fmt = wb.add_format({"bold": True, "font_color": "white", "bg_color": "#1F4E78", "border": 1})
        t_fmt = wb.add_format({"text_wrap": True, "valign": "top", "italic": True, "font_color": "#333333"})

        tri_df.to_excel(writer, sheet_name="Development_Analysis", startrow=1, index=False)
        s1 = writer.sheets["Development_Analysis"]
        s1.write(0, 0, "Cumulative Development Triangle", h_fmt)
        s1.merge_range(len(tri_df) + 3, 0, len(tri_df) + 7, 8, commentary[0] if commentary else "", t_fmt)

        ult_df.to_excel(writer, sheet_name="Reserving_Results", startrow=1, index=False)
        s2 = writer.sheets["Reserving_Results"]
        s2.write(0, 0, "Ultimate Loss Estimates", h_fmt)
        start_ibnr = len(ult_df) + 4
        ibnr_df.to_excel(writer, sheet_name="Reserving_Results", startrow=start_ibnr, index=False)
        s2.write(start_ibnr - 1, 0, "IBNR Estimates", h_fmt)

        if not ave_df.empty:
            ave_df.to_excel(writer, sheet_name="Diagnostics", startrow=1, index=False)
            s3 = writer.sheets["Diagnostics"]
            s3.write(0, 0, "Actual vs Expected Analysis", h_fmt)
            if all(c in ave_df.columns for c in ["Actual", "Expected"]):
                fig, ax = plt.subplots(figsize=(6, 4))
                plot_df = ave_df.iloc[::-1]
                y = range(len(plot_df))
                ax.hlines(y=y, xmin=plot_df["Actual"], xmax=plot_df["Expected"], color="grey")
                ax.scatter(plot_df["Actual"], y, color="#1f77b4")
                ax.scatter(plot_df["Expected"], y, color="#ff7f0e")
                img = _plot_to_buf(fig)
                s3.insert_image(1, 6, "ave_plot.png", {"image_data": img})
            s3.merge_range(len(ave_df) + 3, 0, len(ave_df) + 8, 8, commentary[2] if len(commentary) > 2 else "", t_fmt)

    return "Actuarial_Analysis_Master.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", base64.b64encode(output.getvalue()).decode()


def generate_notebook_context(context_data: dict[str, Any], notes: str):
    txt = f"DATA:\n{json.dumps(context_data, indent=2)}\nNOTES:\n{notes}"
    return "notebook_source.txt", "text/plain", base64.b64encode(txt.encode()).decode()
